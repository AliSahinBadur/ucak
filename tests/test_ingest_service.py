"""End-to-end ingest: parse, clean, chunk, embed and persist a real file.

The fixtures below build genuine PDF and DOCX files on disk (reportlab and
python-docx are both runtime dependencies), so this covers the whole pipeline
`IngestService.ingest` drives rather than a mocked stand-in. Every page carries
well over the OCR character threshold, so the selective-OCR path stays dormant
whether or not the machine running the suite has Tesseract data installed.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from sqlalchemy import func, select

from app.db.models import ChunkEmbedding, Document, DocumentChunk, DocumentPage
from app.services.ingest_service import IngestService
from app.services.search_service import SearchService
from app.services.vector_index import get_vector_index

PAGE_TEXTS = (
    "DAYANIM TEST RAPORU\n"
    "Sasi yorulma olcumleri Torku parkurunda gerceklestirilmistir. "
    "Toplam 12000 kilometre yol verisi toplanmis ve gerinim olcer kanallari "
    "saniyede 200 ornek hizinda kaydedilmistir.",
    "SONUC VE DEGERLENDIRME\n"
    "Kritik bolgelerdeki hasar birikimi tasarim sinir degerinin altinda kalmistir. "
    "Radyator baglanti braketinde gozlenen gerilme yigilmasi icin ek takviye onerilmektedir.",
)


def _write_pdf(path: Path, page_texts: tuple[str, ...] = PAGE_TEXTS) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=A4)
    for text in page_texts:
        vertical = 780
        for line in text.splitlines():
            # reportlab's built-in fonts are Latin-1, and the corpus text is
            # deliberately ASCII-folded Turkish, so no font embedding is needed.
            pdf.drawString(60, vertical, line)
            vertical -= 18
        pdf.showPage()
    pdf.save()
    return path


def _write_docx(path: Path) -> Path:
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading("TITRESIM OLCUM RAPORU", level=1)
    document.add_paragraph(
        "Titresim olcumleri direksiyon simidi ve koltuk rayi uzerinden alinmistir. "
        "Kabin ici gurultu seviyesi 68 dBA olarak olculmustur."
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Kanal"
    table.cell(0, 1).text = "Deger"
    table.cell(1, 0).text = "Koltuk rayi"
    table.cell(1, 1).text = "0.42 m/s2"
    document.save(str(path))
    return path


@pytest.fixture
def storage_dir(tmp_path: Path) -> Path:
    return tmp_path / "documents"


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    return _write_pdf(tmp_path / "2025-BIG-E-DUR-01 Dayanim.pdf")


class ZeroEmbeddingService:
    """Stands in for a provider that produces no signal for this corpus."""

    provider_name = "zero-test"

    def embed_document(self, text: str) -> list[float]:  # noqa: ARG002
        return [0.0, 0.0]

    def embed_query(self, text: str) -> list[float]:  # noqa: ARG002
        return [0.0, 0.0]

    def embed_text(self, text: str) -> list[float]:  # noqa: ARG002
        return [0.0, 0.0]

    @staticmethod
    def has_signal(vector: list[float]) -> bool:
        return any(value != 0.0 for value in vector)

    @staticmethod
    def serialize(vector: list[float]) -> bytes:
        raise AssertionError("serialize must not be reached for a signal-free vector")


def _counts(session, document_id: int) -> tuple[int, int, int]:
    pages = session.scalar(
        select(func.count(DocumentPage.id)).where(DocumentPage.document_id == document_id)
    )
    chunks = session.scalar(
        select(func.count(DocumentChunk.id)).where(DocumentChunk.document_id == document_id)
    )
    embeddings = session.scalar(
        select(func.count(ChunkEmbedding.chunk_id))
        .join(DocumentChunk, DocumentChunk.id == ChunkEmbedding.chunk_id)
        .where(DocumentChunk.document_id == document_id)
    )
    return int(pages), int(chunks), int(embeddings)


# --- the happy path ----------------------------------------------------------


def test_ingest_pdf_persists_pages_chunks_and_embeddings(db_session, sample_pdf, storage_dir) -> None:
    service = IngestService(db_session, storage_dir=storage_dir)

    result = service.ingest(sample_pdf)

    assert result["status"] == "ingested"
    assert result["file_name"] == "2025-BIG-E-DUR-01 Dayanim.pdf"
    assert result["pages"] == 2
    assert result["chunks"] == 2
    assert result["ocr_pages"] == 0
    assert result["embeddings_created"] == 2
    assert result["embedding_provider"] == "token-hash-v1"

    assert _counts(db_session, result["document_id"]) == (2, 2, 2)

    document = db_session.get(Document, result["document_id"])
    assert document.title == "2025-BIG-E-DUR-01 Dayanim"
    assert document.file_type == "pdf"
    assert len(document.file_hash) == 64


def test_ingest_copies_the_source_into_the_storage_directory(db_session, sample_pdf, storage_dir) -> None:
    service = IngestService(db_session, storage_dir=storage_dir)

    result = service.ingest(sample_pdf)

    document = db_session.get(Document, result["document_id"])
    stored = Path(document.file_path)
    assert stored.parent == storage_dir
    assert stored.is_file()
    # Unsafe characters in the original stem become underscores, and the first
    # eight hex digits of the content hash keep distinct files apart.
    assert stored.name == f"2025-BIG-E-DUR-01_Dayanim__{document.file_hash[:8]}.pdf"
    assert sample_pdf.is_file(), "the source file must not be moved"


def test_ingest_keeps_the_page_and_chunk_text_readable(db_session, sample_pdf, storage_dir) -> None:
    result = IngestService(db_session, storage_dir=storage_dir).ingest(sample_pdf)

    pages = db_session.scalars(
        select(DocumentPage)
        .where(DocumentPage.document_id == result["document_id"])
        .order_by(DocumentPage.page_number)
    ).all()
    chunks = db_session.scalars(
        select(DocumentChunk)
        .where(DocumentChunk.document_id == result["document_id"])
        .order_by(DocumentChunk.chunk_order)
    ).all()

    assert [page.page_number for page in pages] == [1, 2]
    assert "yorulma olcumleri" in pages[0].clean_text
    assert pages[0].section_title == "DAYANIM TEST RAPORU"
    assert [chunk.chunk_order for chunk in chunks] == [1, 2]
    assert chunks[0].page_start == 1 and chunks[0].page_end == 1
    assert "Torku parkurunda" in chunks[0].chunk_text


def test_ingested_document_is_immediately_searchable(db_session, sample_pdf, storage_dir) -> None:
    result = IngestService(db_session, storage_dir=storage_dir).ingest(sample_pdf)
    service = SearchService(db_session)

    keyword_hits = service.keyword_search("yorulma olcumleri", limit=5)
    semantic_hits = service.semantic_search("gerinim olcer kanallari", limit=5)

    assert [item["document_id"] for item in keyword_hits] == [result["document_id"]]
    assert [item["document_id"] for item in semantic_hits] == [result["document_id"]]


def test_ingest_refreshes_the_vector_index(db_session, sample_pdf, storage_dir) -> None:
    assert get_vector_index(db_session) is None

    result = IngestService(db_session, storage_dir=storage_dir).ingest(sample_pdf)

    index = get_vector_index(db_session)
    assert index is not None
    assert index.chunk_ids.size == 2
    assert set(index.document_ids.tolist()) == {result["document_id"]}


def test_ingest_docx_reads_headings_paragraphs_and_tables(db_session, tmp_path, storage_dir) -> None:
    docx_path = _write_docx(tmp_path / "titresim.docx")

    result = IngestService(db_session, storage_dir=storage_dir).ingest(docx_path)

    assert result["status"] == "ingested"
    assert result["chunks"] >= 1
    document = db_session.get(Document, result["document_id"])
    assert document.file_type == "docx"
    chunk_text = db_session.scalar(
        select(DocumentChunk.chunk_text).where(DocumentChunk.document_id == document.id)
    )
    assert "Kabin ici gurultu" in chunk_text
    assert "Koltuk rayi | 0.42 m/s2" in chunk_text


# --- the duplicate path ------------------------------------------------------


def test_re_ingesting_the_same_file_is_a_no_op_duplicate(db_session, sample_pdf, storage_dir) -> None:
    service = IngestService(db_session, storage_dir=storage_dir)
    first = service.ingest(sample_pdf)

    second = service.ingest(sample_pdf)

    assert second["status"] == "duplicate"
    assert second["document_id"] == first["document_id"]
    assert second["embeddings_created"] == 0
    assert "pages" not in second
    assert db_session.scalar(select(func.count(Document.id))) == 1
    assert _counts(db_session, first["document_id"]) == (2, 2, 2)


def test_duplicates_are_detected_by_content_not_by_file_name(db_session, tmp_path, storage_dir) -> None:
    service = IngestService(db_session, storage_dir=storage_dir)
    original = _write_pdf(tmp_path / "ilk-adi.pdf")
    # A byte-for-byte copy, because a freshly generated PDF carries its own
    # creation timestamp and would hash differently.
    renamed_copy = tmp_path / "bambaska-bir-ad.pdf"
    renamed_copy.write_bytes(original.read_bytes())
    first = service.ingest(original)

    second = service.ingest(renamed_copy)

    assert second["status"] == "duplicate"
    assert second["document_id"] == first["document_id"]
    # The stored file keeps the name it was first ingested under.
    assert second["file_name"] == "ilk-adi.pdf"


def test_different_content_is_not_a_duplicate(db_session, tmp_path, storage_dir) -> None:
    service = IngestService(db_session, storage_dir=storage_dir)
    first = service.ingest(_write_pdf(tmp_path / "a.pdf"))

    second = service.ingest(_write_pdf(tmp_path / "b.pdf", page_texts=(PAGE_TEXTS[1],)))

    assert second["status"] == "ingested"
    assert second["document_id"] != first["document_id"]
    assert db_session.scalar(select(func.count(Document.id))) == 2


def test_original_file_name_overrides_the_temp_upload_name(db_session, tmp_path, storage_dir) -> None:
    upload_temp = _write_pdf(tmp_path / "tmp8f21a.pdf")

    result = IngestService(db_session, storage_dir=storage_dir).ingest(
        upload_temp, original_file_name="Ağır Ticari Dayanım Raporu.pdf"
    )

    document = db_session.get(Document, result["document_id"])
    assert document.file_name == "Ağır Ticari Dayanım Raporu.pdf"
    assert document.title == "Ağır Ticari Dayanım Raporu"
    # Spaces collapse to underscores; Turkish letters are alphanumeric to
    # str.isalnum(), so they survive the on-disk name unchanged.
    assert Path(document.file_path).name == f"Ağır_Ticari_Dayanım_Raporu__{document.file_hash[:8]}.pdf"


# --- failure and degradation paths ------------------------------------------


def test_ingest_rejects_an_unsupported_extension(db_session, tmp_path, storage_dir) -> None:
    text_file = tmp_path / "notlar.txt"
    text_file.write_text("merhaba", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        IngestService(db_session, storage_dir=storage_dir).ingest(text_file)

    assert db_session.scalar(select(func.count(Document.id))) == 0


def test_chunks_without_embedding_signal_are_still_stored(db_session, sample_pdf, storage_dir) -> None:
    service = IngestService(
        db_session, storage_dir=storage_dir, embedding_service=ZeroEmbeddingService()
    )

    result = service.ingest(sample_pdf)

    assert result["status"] == "ingested"
    assert result["chunks"] == 2
    assert result["embeddings_created"] == 0
    assert result["embedding_provider"] == "zero-test"
    assert _counts(db_session, result["document_id"]) == (2, 2, 0)
    # No embeddings means no semantic index, but keyword search still works.
    assert get_vector_index(db_session) is None
    assert SearchService(db_session).keyword_search("yorulma", limit=5)


def test_ingest_creates_its_storage_directory_on_demand(db_session, tmp_path) -> None:
    nested = tmp_path / "deeply" / "nested" / "documents"
    assert not nested.exists()

    IngestService(db_session, storage_dir=nested)

    assert nested.is_dir()
