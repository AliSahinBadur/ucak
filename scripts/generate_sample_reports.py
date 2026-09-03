"""Generate a synthetic CAE report corpus for hands-on testing.

Every report here is invented: fictional vehicles (SYN-Bus 12, SYN-Van 3),
fictional people, and report codes prefixed `2026-SYN-` so nothing can be
mistaken for a real company document. The numbers are plausible but made up.

The corpus is built so that each document exercises a known part of the review
engine -- some pass cleanly, some carry exactly one planted defect. Run
`scripts/verify_sample_reports.py` after generating to see the findings each
one produces.

    python scripts/generate_sample_reports.py            # -> ./sample_reports
    python scripts/generate_sample_reports.py --out D:\\demo

Requires reportlab and python-docx, both already in requirements.txt.
"""

from __future__ import annotations

import argparse
import csv
from pathlib import Path
import shutil
import sys

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


BASE_DIR = Path(__file__).resolve().parent.parent
DEFAULT_OUT = BASE_DIR / "sample_reports"

PAGE_W, PAGE_H = A4
MARGIN_X = 56
TOP_Y = PAGE_H - 68
LEADING = 15.5
BODY_SIZE = 10.5


def _register_fonts() -> tuple[str, str]:
    """Use a real TTF so Turkish characters survive into the PDF text layer.

    Helvetica's WinAnsi encoding has no glyph for s-cedilla, g-breve or
    dotless i, so a report written with it would come back out of the parser
    with holes -- exactly the thing this corpus exists to test.
    """
    candidates = [
        (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")),
        (Path("C:/Windows/Fonts/segoeui.ttf"), Path("C:/Windows/Fonts/segoeuib.ttf")),
        (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
         Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")),
    ]
    for regular, bold in candidates:
        if regular.exists():
            pdfmetrics.registerFont(TTFont("Body", str(regular)))
            if bold.exists():
                pdfmetrics.registerFont(TTFont("BodyBold", str(bold)))
                return "Body", "BodyBold"
            return "Body", "Body"
    print("! No Unicode TTF found; falling back to Helvetica (Turkish glyphs will be lost).")
    return "Helvetica", "Helvetica-Bold"


BODY_FONT, BOLD_FONT = "Helvetica", "Helvetica-Bold"

# Lines rendered in bold: headings and cover labels read as headings in a real report.
BOLD_PREFIXES = (
    "KAPSAM", "SONUÇLAR", "SONUCLAR", "YÖNTEM", "YONTEM", "GİRİŞ", "GIRIS",
    "MODEL", "MALZEME", "TEST DÜZENİ", "TEST DUZENI", "DEĞERLENDİRME",
    "ÖLÇÜM DÜZENİ", "ANALİZ", "REVİZYON", "SINIR ŞARTLARI", "PROSEDÜR",
)


def _is_bold(line: str) -> bool:
    return any(line.startswith(prefix) for prefix in BOLD_PREFIXES)


def write_pdf(path: Path, pages: list[list[str]]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.setTitle(path.stem)
    for page_lines in pages:
        y = TOP_Y
        for line in page_lines:
            if line == "@@BLANK_PAGE@@":
                # A page with no text layer at all: a scanned drawing, as far as
                # the parser is concerned. This is what extraction.no_text is for.
                pdf.setStrokeColorRGB(0.35, 0.35, 0.35)
                pdf.setLineWidth(1.2)
                pdf.rect(MARGIN_X, PAGE_H / 2 - 130, PAGE_W - 2 * MARGIN_X, 260)
                for step in range(1, 9):
                    x = MARGIN_X + step * (PAGE_W - 2 * MARGIN_X) / 9
                    pdf.line(x, PAGE_H / 2 - 130, x, PAGE_H / 2 + 130)
                pdf.line(MARGIN_X, PAGE_H / 2, PAGE_W - MARGIN_X, PAGE_H / 2)
                continue
            pdf.setFont(BOLD_FONT if _is_bold(line) else BODY_FONT, BODY_SIZE)
            pdf.drawString(MARGIN_X, y, line)
            y -= LEADING
        pdf.showPage()
    pdf.save()


def write_docx(path: Path, pages: list[list[str]]) -> None:
    from docx import Document as DocxDocument

    document = DocxDocument()
    for index, page_lines in enumerate(pages):
        if index:
            document.add_page_break()
        for line in page_lines:
            if not line.strip():
                document.add_paragraph("")
            elif _is_bold(line):
                document.add_heading(line, level=2)
            else:
                document.add_paragraph(line)
    document.save(str(path))


# --------------------------------------------------------------------------
# The corpus. `expect` documents the planted defect so the guide and the
# verification script agree on what each file is for.
# --------------------------------------------------------------------------

COVER = "SYN Otomotiv - Mühendislik Raporu"


def cover_block(code: str, date: str, author: str, checker: str, vehicle: str,
                date_label: str = "TARIH") -> list[str]:
    return [
        COVER,
        "",
        f"RAPOR NO: {code}",
        f"{date_label}: {date}",
        f"HAZIRLAYAN: {author}",
        f"KONTROL: {checker}",
        f"ARAÇ: {vehicle}",
        "",
    ]


REPORTS: list[dict] = [
    # ---------------- clean, one per discipline ----------------
    {
        "code": "2026-SYN-e-NVH-01",
        "format": "pdf",
        "vehicle": "SYN-Bus 12",
        "title": "Sürücü koltuğu titreşim konforu ölçümü",
        "discipline": "NVH",
        "date": "2026-03-14",
        "author": "A. Yılmaz",
        "checker": "M. Demir",
        "expect": "clean - all NVH profile rules pass",
        "pages": [
            [
                *cover_block("2026-SYN-e-NVH-01", "2026-03-14", "A. Yılmaz", "M. Demir", "SYN-Bus 12"),
                "KAPSAM",
                "Bu raporda SYN-Bus 12 aracının sürücü koltuğu titreşim konforu",
                "değerlendirilmiştir. Ölçümler koltuk üzerinden yerleştirilen üç eksenli",
                "ivmeölçer ile x ekseni, y ekseni ve z ekseni yönlerinde alınmıştır.",
                "Ölçümler 50 km/h sabit hızda bozuk yol parkurunda ve 90 km/h otoban",
                "parkurunda tekrarlanmıştır. Motor rölanti durumunda referans ölçümü",
                "ayrıca kaydedilmiştir. Ölçüm noktası koltuk minderi bağlantı",
                "noktasıdır ve sensör üç eksenli olarak sabitlenmiştir.",
            ],
            [
                "YÖNTEM",
                "Sinyal işleme ISO 2631-1 standardına göre yapılmıştır. Ham ivme",
                "sinyalleri 0-80 Hz frekans aralığında bant geçiren filtre ile",
                "işlenmiş ve frekans ağırlıklandırması uygulanmıştır. Örnekleme",
                "frekansı 2048 Hz olarak seçilmiştir. Her ölçüm için gRMS ve crest",
                "faktör değerleri hesaplanmıştır.",
                "",
                "Tablo 1 - Parkur bazında gRMS değerleri",
                "Bozuk yol 50 km/h ölçümünde 0,42 m/s2 değeri kaydedilmiştir.",
                "Otoban 90 km/h ölçümünde 0,18 m/s2 değeri kaydedilmiştir.",
                "Rölanti ölçümünde 0,05 m/s2 değeri kaydedilmiştir.",
                "",
                "Tablo 1'de verilen değerler üç tekrarın ortalamasıdır.",
            ],
            [
                "SONUÇLAR",
                "ISO 2631-1 kabul kriteri olan 0,50 m/s2 sınır değeri hiçbir parkurda",
                "aşılmamıştır. En yüksek değer bozuk yol parkurunda 0,42 m/s2 olarak",
                "ölçülmüştür. Sonuç uygun olarak değerlendirilmiştir ve ek bir",
                "iyileştirme çalışması önerilmemiştir.",
                "",
                "Şekil 1 - Bozuk yol parkuru frekans spektrumu",
                "Şekil 1'de baskın bileşenin 12 Hz civarında olduğu görülmektedir.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-CFD-01",
        "format": "pdf",
        "vehicle": "SYN-Bus 12",
        "title": "Yolcu kabini defrost analizi",
        "discipline": "CFD",
        "date": "2026-02-27",
        "author": "E. Kaya",
        "checker": "M. Demir",
        "expect": "clean - all CFD profile rules pass",
        "pages": [
            [
                *cover_block("2026-SYN-e-CFD-01", "2026-02-27", "E. Kaya", "M. Demir", "SYN-Bus 12"),
                "KAPSAM",
                "SYN-Bus 12 ön cam defrost performansı hesaplamalı akışkanlar",
                "dinamiği ile incelenmiştir. Çözüm Fluent solver üzerinde k-epsilon",
                "türbülans modeli ile yürütülmüştür. Analiz kabin içi hava dağılımını",
                "ve cam yüzeyindeki sıcaklık artışını kapsamaktadır.",
            ],
            [
                "SINIR ŞARTLARI",
                "Üfleyici çıkışı inlet sınır şartı olarak tanımlanmış, giriş hızı",
                "8,50 m/s ve giriş sıcaklığı 45,0 derece C alınmıştır. Kabin çıkışı",
                "outlet basınç sınır şartı ile modellenmiştir. Toplam debi 0,108 m3/s",
                "olarak ayarlanmıştır. Dış ortam sıcaklığı -10,0 derece C kabul",
                "edilmiştir.",
                "",
                "MODEL",
                "Ağ yapısı poliheral mesh olarak oluşturulmuş, toplam hücre sayısı",
                "4,20 milyon olarak elde edilmiştir. Cam yüzeyinde inflation katmanı",
                "kullanılmış ve y+ değeri 1,20 civarında tutulmuştur. Yakınsama",
                "residual değerleri 1e-4 seviyesinin altına indiğinde kabul edilmiş,",
                "iterasyon sayısı 1200 olarak kaydedilmiştir.",
                "",
                "Tablo 1 - Cam bölgelerine göre defrost süreleri",
                "Tablo 1'de A bölgesi için 11 dakika, B bölgesi için 17 dakika",
                "değerleri verilmiştir.",
            ],
            [
                "SONUÇLAR",
                "Hesaplanan defrost süreleri hedef değer olan 20 dakikanın altında",
                "kalmıştır. Mevcut tasarım referans tasarım ile karşılaştırıldığında",
                "A bölgesinde 3 dakikalık iyileşme elde edilmiştir. Sonuç kabul",
                "kriterine göre uygun bulunmuştur.",
                "",
                "Şekil 1 - Cam yüzeyi sıcaklık dağılımı",
                "Şekil 1'de sıcaklık dağılımının simetrik olduğu görülmektedir.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-DUR-01",
        "format": "pdf",
        "vehicle": "SYN-Van 3",
        "title": "Batarya taşıyıcı braket statik dayanım analizi",
        "discipline": "DURABILITY",
        "date": "2026-01-19",
        "author": "S. Aydın",
        "checker": "A. Yılmaz",
        "expect": "clean - all durability profile rules pass",
        "pages": [
            [
                *cover_block("2026-SYN-e-DUR-01", "2026-01-19", "S. Aydın", "A. Yılmaz", "SYN-Van 3"),
                "KAPSAM",
                "SYN-Van 3 batarya taşıyıcı braketinin statik dayanımı sonlu eleman",
                "yöntemi ile incelenmiştir. Analiz yol yükleri altında braket ve",
                "bağlantı elemanlarının gerilme seviyelerini kapsamaktadır.",
                "",
                "MALZEME",
                "Braket malzemesi S355 yapı çeliğidir. Hesapta elastisite modülü",
                "210000 MPa, poisson oranı 0,30 ve akma dayanımı 355 MPa olarak",
                "kullanılmıştır.",
            ],
            [
                "MODEL",
                "Sonlu eleman modeli 3D eleman ile oluşturulmuş, ortalama eleman",
                "boyutu 4,00 mm seçilmiştir. Braket gövdesinde hexa eleman, kaynak",
                "bölgelerinde tetra eleman kullanılmıştır. Cıvata bağlantıları temas",
                "tanımı ile modellenmiş, kaynak dikişleri sürekli bağlantı olarak",
                "kabul edilmiştir.",
                "",
                "SINIR ŞARTLARI",
                "Braket şase bağlantı noktalarından ankastre olarak sabitlenmiştir.",
                "Uygulanan yük batarya kütlesi 320 kg üzerinden düşey doğrultuda",
                "3,00 g ivme ile tanımlanmıştır. Yanal doğrultuda 1,50 g ivme ayrıca",
                "uygulanmıştır. Mesnet noktalarında dönme serbestliği kısıtlanmıştır.",
                "",
                "Tablo 1 - Yük durumlarına göre maksimum gerilme",
                "Tablo 1'de düşey yük için 208 MPa, yanal yük için 141 MPa değerleri",
                "listelenmiştir.",
            ],
            [
                "SONUÇLAR",
                "Hesaplanan maksimum von Mises gerilmesi 208 MPa olarak bulunmuştur.",
                "Malzemenin akma dayanımı 355 MPa olduğundan emniyet katsayısı 1,70",
                "olarak hesaplanmıştır. Kabul kriteri olan 1,50 emniyet katsayısı",
                "sağlanmıştır ve tasarım emniyetli olarak değerlendirilmiştir.",
                "Maksimum deformasyon 1,80 mm seviyesindedir.",
                "",
                "Şekil 1 - von Mises gerilme dağılımı",
                "Şekil 1'de en yüksek gerilmenin kaynak geçiş bölgesinde oluştuğu",
                "görülmektedir.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-TEST-01",
        "format": "docx",
        "vehicle": "SYN-Van 3",
        "title": "Soğutma sistemi dayanıklılık testi",
        "discipline": "TEST",
        "date": "2026-04-02",
        "author": "B. Şahin",
        "checker": "E. Kaya",
        "expect": "clean - all test profile rules pass; also the DOCX parser path",
        "pages": [
            [
                *cover_block("2026-SYN-e-TEST-01", "2026-04-02", "B. Şahin", "E. Kaya", "SYN-Van 3"),
                "KAPSAM",
                "SYN-Van 3 prototip aracının soğutma sistemi dayanıklılık testi",
                "raporlanmıştır. Test objesi seri no SV3-0042 olan prototip araçtır",
                "ve inverter soğutma devresi konfigürasyonu ile test edilmiştir.",
            ],
            [
                "TEST DÜZENİ",
                "Ölçüm cihazı olarak dört kanallı termokupl seti ve bir data logger",
                "kullanılmıştır. Devre giriş ve çıkışına manometre bağlanmıştır.",
                "Cihaz seri no TL-2291 olup kalibrasyon sertifika no KAL-2025-338",
                "ile geçerlidir. Ortam sıcaklığı 38,0 derece C olarak sabitlenmiştir.",
                "",
                "PROSEDÜR",
                "Test yöntemi üç aşamalı olarak uygulanmıştır. Araç 120 dakika",
                "boyunca 90 km/h sabit hızda, ardından 60 dakika boyunca 1500 rpm",
                "motor devrinde tam yük altında çalıştırılmıştır. Son aşamada 30",
                "dakika rölanti soğuma çevrimi uygulanmıştır. Toplam test süresi",
                "210 dakikadır.",
                "",
                "Tablo 1 - Aşamalara göre maksimum soğutucu sıcaklığı",
                "Tablo 1'de birinci aşama için 82,0 derece C, ikinci aşama için",
                "94,0 derece C değerleri verilmiştir.",
            ],
            [
                "SONUÇLAR",
                "Kabul kriteri maksimum 105,0 derece C soğutucu sıcaklığıdır.",
                "Ölçülen en yüksek değer 94,0 derece C olarak kaydedilmiştir.",
                "Test sonucu OK olarak değerlendirilmiş ve sistem uygun bulunmuştur.",
                "",
                "Şekil 1 - Test süresince sıcaklık eğrisi",
                "Şekil 1'de ikinci aşamada plato oluştuğu görülmektedir.",
            ],
        ],
    },

    # ---------------- one planted profile defect each ----------------
    {
        "code": "2026-SYN-e-NVH-02",
        "format": "pdf",
        "vehicle": "SYN-Bus 12",
        "title": "Ayna titreşimi ön inceleme",
        "discipline": "NVH",
        "date": "2026-03-21",
        "author": "A. Yılmaz",
        "checker": "M. Demir",
        "expect": "nvh.measurement_setup -> needs_review (measurement axis never stated)",
        "pages": [
            [
                *cover_block("2026-SYN-e-NVH-02", "2026-03-21", "A. Yılmaz", "M. Demir", "SYN-Bus 12"),
                "KAPSAM",
                "Dış dikiz aynasındaki titreşim şikayeti incelenmiştir. Ölçüm ayna",
                "üzerinden yerleştirilen sensör ile 70 km/h otoban parkurunda",
                "alınmıştır. Tek kanal kaydedilmiş, kanal bilgisi ekte verilmiştir.",
            ],
            [
                "YÖNTEM",
                "Ham sinyal 0-200 Hz frekans aralığında filtre ile işlenmiş, gRMS",
                "ve peak değerleri hesaplanmıştır. Örnekleme frekansı 4096 Hz",
                "olarak seçilmiştir.",
                "",
                "SONUÇLAR",
                "ISO 2631-1 sınır değeri ile karşılaştırıldığında ölçülen 0,31 m/s2",
                "değeri kabul kriterinin altındadır. Sonuç uygun bulunmuştur.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-CFD-02",
        "format": "pdf",
        "vehicle": "SYN-Van 3",
        "title": "Motor bölmesi hava akışı ön çalışması",
        "discipline": "CFD",
        "date": "2026-02-11",
        "author": "E. Kaya",
        "checker": "S. Aydın",
        "expect": "cfd.numerical_evidence -> needs_review (mesh stated, convergence never shown)",
        "pages": [
            [
                *cover_block("2026-SYN-e-CFD-02", "2026-02-11", "E. Kaya", "S. Aydın", "SYN-Van 3"),
                "KAPSAM",
                "SYN-Van 3 motor bölmesi hava akışı incelenmiştir. Çözüm Fluent",
                "solver üzerinde k-omega türbülans modeli ile yürütülmüştür.",
                "",
                "SINIR ŞARTLARI",
                "Radyatör önü inlet olarak tanımlanmış, giriş hızı 12,0 m/s",
                "alınmıştır. Bölme çıkışı outlet basınç sınır şartı ile",
                "modellenmiştir. Fan devri 2400 rpm olarak ayarlanmıştır.",
            ],
            [
                "MODEL",
                "Ağ yapısı tetra mesh olarak oluşturulmuş, eleman sayısı 6,80",
                "milyon olarak elde edilmiştir. Ağ kalitesi kontrol edilmiştir.",
                "",
                "SONUÇLAR",
                "Radyatör yüzeyinden geçen debi 0,240 m3/s olarak hesaplanmıştır.",
                "Mevcut tasarım hedef debi ile karşılaştırılmış ve kabul kriterine",
                "göre uygun bulunmuştur.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-DUR-02",
        "format": "pdf",
        "vehicle": "SYN-Van 3",
        "title": "Çekme kancası bağlantı analizi",
        "discipline": "DURABILITY",
        "date": "2026-01-30",
        "author": "S. Aydın",
        "checker": "A. Yılmaz",
        "expect": "durability.result_criterion -> needs_review (a verdict with no acceptance basis)",
        "pages": [
            [
                *cover_block("2026-SYN-e-DUR-02", "2026-01-30", "S. Aydın", "A. Yılmaz", "SYN-Van 3"),
                "KAPSAM",
                "Çekme kancası bağlantı bölgesi statik olarak incelenmiştir.",
                "",
                "MALZEME",
                "Kanca malzemesi S235 yapı çeliğidir. Elastisite modülü 210000 MPa",
                "olarak alınmıştır.",
            ],
            [
                "MODEL",
                "Sonlu eleman modeli 3D eleman ile kurulmuş, eleman boyutu 5,00 mm",
                "seçilmiştir. Cıvata bağlantıları temas tanımı ile modellenmiştir.",
                "",
                "SINIR ŞARTLARI",
                "Şase bağlantı yüzeyleri sabitlenmiş, kancaya 25,0 kN çekme kuvveti",
                "uygulanmıştır. Mesnet bölgesinde ankastre kabul yapılmıştır.",
                "",
                "SONUÇLAR",
                "Maksimum von Mises gerilmesi 187 MPa olarak hesaplanmıştır.",
                "Maksimum deformasyon 0,90 mm seviyesindedir. Sonuçlar tasarım",
                "ekibine iletilmiştir.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-TEST-02",
        "format": "pdf",
        "vehicle": "SYN-Bus 12",
        "title": "Kapı mekanizması çevrim testi",
        "discipline": "TEST",
        "date": "2026-04-18",
        "author": "B. Şahin",
        "checker": "E. Kaya",
        "expect": "test.measurement_traceability -> needs_review at INFO severity (no calibration record)",
        "pages": [
            [
                *cover_block("2026-SYN-e-TEST-02", "2026-04-18", "B. Şahin", "E. Kaya", "SYN-Bus 12"),
                "KAPSAM",
                "SYN-Bus 12 orta kapı mekanizmasının çevrim dayanımı test",
                "edilmiştir. Test objesi seri no SB12-0117 olan prototip kapı",
                "modülüdür.",
            ],
            [
                "TEST DÜZENİ",
                "Ölçüm cihazı olarak kuvvet sensörü ve çevrim sayacı kullanılmıştır.",
                "Ortam sıcaklığı 23,0 derece C olarak tutulmuştur.",
                "",
                "PROSEDÜR",
                "Test yöntemi olarak sürekli açma kapama çevrimi uygulanmıştır.",
                "Toplam 20000 çevrim, 180 dakika süre boyunca yürütülmüştür.",
                "",
                "SONUÇLAR",
                "Kabul kriteri 15000 çevrim sonunda fonksiyon kaybı olmamasıdır.",
                "Test sonucu OK olarak değerlendirilmiş, mekanizma uygun",
                "bulunmuştur.",
            ],
        ],
    },

    # ---------------- general-rule defects ----------------
    {
        "code": "2026-SYN-e-GEN-01",
        "format": "pdf",
        "vehicle": "SYN-Bus 12",
        "title": "Ön süspansiyon notu",
        "discipline": "DURABILITY",
        "date": "2026-05-06",
        "author": "",
        "checker": "",
        "expect": "metadata.required_fields + structure.required_sections (no cover fields, no SONUÇLAR)",
        "pages": [
            [
                COVER,
                "",
                "Ön süspansiyon salıncak kolu ön inceleme notu",
                "SYN-Bus 12",
                "",
                "KAPSAM",
                "Salıncak kolu bağlantı bölgesinde gözlemlenen aşınma üzerine kısa",
                "bir inceleme yapılmıştır. Malzeme S355 olarak teyit edilmiştir.",
                "Uygulanan yük 18,0 kN olarak alınmış, mesnet noktaları",
                "sabitlenmiştir. Sonlu eleman modeli 3D eleman ile kurulmuş, temas",
                "tanımı kullanılmıştır. Gerilme dağılımı incelenmiş, akma dayanımı",
                "ile karşılaştırılmış ve emniyet katsayısı hesaplanmıştır.",
                "Bu not resmi rapor yerine geçmez; ayrıntılı değerlendirme",
                "ayrı bir belgede raporlanacaktır.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-GEN-02",
        "format": "pdf",
        "vehicle": "SYN-Van 3",
        "title": "Aks yükü ölçüm raporu",
        "discipline": "TEST",
        "date": "2026-05-12",
        "author": "B. Şahin",
        "checker": "M. Demir",
        "expect": "captions.sequence (Tablo 2 missing) + captions.title (Şekil 2 untitled) "
                  "+ captions.references (Tablo 5 referenced but absent)",
        "pages": [
            [
                *cover_block("2026-SYN-e-GEN-02", "2026-05-12", "B. Şahin", "M. Demir", "SYN-Van 3"),
                "KAPSAM",
                "SYN-Van 3 aks yükü dağılımı tartım platformunda ölçülmüştür.",
                "Test objesi seri no SV3-0051 prototip araçtır. Ölçüm cihazı olarak",
                "kalibrasyonlu tartım platformu kullanılmış, cihaz seri no TP-8842",
                "ve kalibrasyon sertifika no KAL-2026-014 kaydedilmiştir. Ortam",
                "sıcaklığı 21,0 derece C olarak tutulmuştur.",
            ],
            [
                "PROSEDÜR",
                "Test yöntemi olarak boş ve tam yüklü iki durum uygulanmıştır.",
                "Her durumda 30 dakika bekleme süresi sonunda ölçüm alınmıştır.",
                "",
                "Tablo 1 - Boş araç aks yükleri",
                "Tablo 1'de ön aks 1240 kg, arka aks 1580 kg olarak verilmiştir.",
                "",
                "Tablo 3 - Tam yüklü araç aks yükleri",
                "Tablo 3'te ön aks 1390 kg, arka aks 2410 kg olarak verilmiştir.",
                "",
                "Ayrıntılı dağılım Tablo 5'te sunulmuştur.",
            ],
            [
                "SONUÇLAR",
                "Kabul kriteri olan 2600 kg arka aks sınırı aşılmamıştır. Test",
                "sonucu OK olarak değerlendirilmiş ve araç uygun bulunmuştur.",
                "",
                "Şekil 1 - Boş araç yük dağılımı",
                "Şekil 1'de dağılımın dengeli olduğu görülmektedir.",
                "",
                "Şekil 2",
                "Şekil 2'de tam yüklü durum gösterilmektedir.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-GEN-03",
        "format": "pdf",
        "vehicle": "SYN-Bus 12",
        "title": "Klima performans özeti",
        "discipline": "CFD",
        "date": "2026-05-20",
        "author": "E. Kaya",
        "checker": "S. Aydın",
        "expect": "numbers.decimal_style (comma and dot mixed) + content.embedded_paths (a Windows path in the text)",
        "pages": [
            [
                *cover_block("2026-SYN-e-GEN-03", "2026-05-20", "E. Kaya", "S. Aydın", "SYN-Bus 12"),
                "KAPSAM",
                "Klima ünitesi hava debisi özeti. Çözüm Fluent solver ve k-epsilon",
                "modeli ile yapılmış, inlet ve outlet sınır şartları tanımlanmıştır.",
                "Ağ yapısı mesh hücre sayısı 3,10 milyon olarak kurulmuş, residual",
                "yakınsama izlenmiştir.",
            ],
            [
                "SONUÇLAR",
                "Ön bölge çıkış debisi 0,108 m3/s olarak hesaplanmıştır.",
                "Arka bölge çıkış debisi 0.245 m3/s olarak hesaplanmıştır.",
                "Kabin ortalama hava hızı 1.85 m/s seviyesindedir.",
                "Mevcut tasarım hedef değer ile karşılaştırılmış ve uygun",
                "bulunmuştur.",
                "",
                "Ham çözüm dosyaları C:\\CAE\\2026\\SYN-Bus12\\klima\\case_final.cas",
                "yolunda saklanmaktadır.",
                "",
                "Şekil 1 - Kabin hava hızı dağılımı",
                "Şekil 1'de ön bölgenin daha yüksek hızda olduğu görülmektedir.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-GEN-04",
        "format": "pdf",
        "vehicle": "SYN-Van 3",
        "title": "Şase tarama eki",
        "discipline": "DURABILITY",
        "date": "2026-05-25",
        "author": "S. Aydın",
        "checker": "A. Yılmaz",
        "expect": "extraction.no_text -> status FAIL (page 2 has a drawing and no text layer)",
        "pages": [
            [
                *cover_block("2026-SYN-e-GEN-04", "2026-05-25", "S. Aydın", "A. Yılmaz", "SYN-Van 3"),
                "KAPSAM",
                "Şase tarama çıktısı ektedir. Malzeme S355, uygulanan yük 22,0 kN,",
                "mesnet noktaları sabitlenmiştir. Sonlu eleman modeli 3D eleman ve",
                "temas tanımı ile kurulmuştur.",
                "",
                "SONUÇLAR",
                "Gerilme sonucu akma dayanımı ile karşılaştırılmış, emniyet",
                "katsayısı kabul kriterini sağlamıştır.",
            ],
            ["@@BLANK_PAGE@@"],
        ],
    },
    {
        "code": "2026-SYN-e-GEN-05",
        "format": "pdf",
        "vehicle": "SYN-Bus 12",
        "title": "Fren balata sıcaklık ölçümü",
        "discipline": "TEST",
        "date": "2026-06-03",
        "author": "B. Şahin",
        "checker": "E. Kaya",
        "expect": "metadata.required_fields -> the TARIH miss caused by the Turkish capital I bug "
                  "(cover says TARİH with U+0130)",
        "pages": [
            [
                *cover_block("2026-SYN-e-GEN-05", "2026-06-03", "B. Şahin", "E. Kaya",
                             "SYN-Bus 12", date_label="TARİH"),
                "KAPSAM",
                "Fren balata sıcaklıkları iniş parkurunda ölçülmüştür. Test objesi",
                "seri no SB12-0120 prototip araçtır. Ölçüm cihazı termokupl seti,",
                "cihaz seri no TK-4410, kalibrasyon sertifika no KAL-2026-051.",
                "Ortam sıcaklığı 26,0 derece C olarak kaydedilmiştir.",
            ],
            [
                "PROSEDÜR",
                "Test yöntemi olarak 8 km iniş parkurunda 60 km/h sabit hız",
                "uygulanmıştır. Toplam süre 45 dakikadır.",
                "",
                "SONUÇLAR",
                "Kabul kriteri maksimum 350,0 derece C balata sıcaklığıdır.",
                "Ölçülen en yüksek değer 288,0 derece C olmuştur. Test sonucu OK",
                "olarak değerlendirilmiş, sistem uygun bulunmuştur.",
            ],
        ],
    },

    # ---------------- revision pair ----------------
    {
        "code": "2026-SYN-e-DUR-03-RevA",
        "format": "pdf",
        "vehicle": "SYN-Van 3",
        "title": "Koltuk bağlantısı dayanım analizi Rev A",
        "discipline": "DURABILITY",
        "date": "2026-02-05",
        "author": "S. Aydın",
        "checker": "A. Yılmaz",
        "expect": "revision pair, older side: material group missing -> durability.material_definition",
        "pages": [
            [
                *cover_block("2026-SYN-e-DUR-03-RevA", "2026-02-05", "S. Aydın", "A. Yılmaz", "SYN-Van 3"),
                "KAPSAM",
                "Koltuk bağlantı braketi dayanım analizi ilk revizyonudur.",
                "",
                "SINIR ŞARTLARI",
                "Braket taban yüzeyinden sabitlenmiş, uygulanan yük 12,0 kN olarak",
                "tanımlanmıştır. Mesnet bölgesi ankastre kabul edilmiştir.",
            ],
            [
                "MODEL",
                "Sonlu eleman modeli 3D eleman ile kurulmuş, eleman boyutu 3,00 mm",
                "seçilmiştir. Cıvata bağlantıları temas tanımı ile modellenmiştir.",
                "",
                "SONUÇLAR",
                "Maksimum von Mises gerilmesi 268 MPa bulunmuştur. Kabul kriteri",
                "olan emniyet katsayısı 1,50 değerinin altında kalınmıştır ve",
                "tasarım emniyetsiz olarak değerlendirilmiştir. Revizyon",
                "önerilmektedir.",
            ],
        ],
    },
    {
        "code": "2026-SYN-e-DUR-03-RevB",
        "format": "pdf",
        "vehicle": "SYN-Van 3",
        "title": "Koltuk bağlantısı dayanım analizi Rev B",
        "discipline": "DURABILITY",
        "date": "2026-03-08",
        "author": "S. Aydın",
        "checker": "A. Yılmaz",
        "expect": "revision pair, newer side: material group added, so the Rev A finding is resolved",
        "pages": [
            [
                *cover_block("2026-SYN-e-DUR-03-RevB", "2026-03-08", "S. Aydın", "A. Yılmaz", "SYN-Van 3"),
                "KAPSAM",
                "Koltuk bağlantı braketi dayanım analizi ikinci revizyonudur.",
                "Braket kesiti kalınlaştırılmış ve yeniden hesaplanmıştır.",
                "",
                "MALZEME",
                "Braket malzemesi S355 yapı çeliğidir. Elastisite modülü 210000 MPa,",
                "poisson oranı 0,30 ve akma dayanımı 355 MPa olarak alınmıştır.",
            ],
            [
                "SINIR ŞARTLARI",
                "Braket taban yüzeyinden sabitlenmiş, uygulanan yük 12,0 kN olarak",
                "tanımlanmıştır. Mesnet bölgesi ankastre kabul edilmiştir.",
                "",
                "MODEL",
                "Sonlu eleman modeli 3D eleman ile kurulmuş, eleman boyutu 3,00 mm",
                "seçilmiştir. Cıvata bağlantıları temas tanımı ile modellenmiştir.",
                "",
                "SONUÇLAR",
                "Maksimum von Mises gerilmesi 196 MPa bulunmuştur. Akma dayanımı",
                "355 MPa üzerinden emniyet katsayısı 1,81 olarak hesaplanmıştır.",
                "Kabul kriteri olan 1,50 değeri sağlanmış, tasarım emniyetli olarak",
                "değerlendirilmiştir.",
            ],
        ],
    },
]


def build_catalog_rows() -> list[list[str]]:
    """Six positional columns: code, vehicle, title, discipline, date, authors.

    The importer reads by position, not by header name, and skips anything that
    looks like a header row -- so the header line below is written for the human
    opening the file and ignored by the parser.
    """
    rows = [["Rapor Kodu", "Arac", "Baslik", "Disiplin", "Tarih", "Hazirlayan"]]
    for report in REPORTS:
        rows.append([
            report["code"],
            report["vehicle"],
            report["title"],
            report["discipline"],
            report["date"],
            report["author"] or "Belirtilmemis",
        ])
    return rows


def main() -> int:
    global BODY_FONT, BOLD_FONT

    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT, help="output folder (default: ./sample_reports)")
    parser.add_argument("--clean", action="store_true", help="empty the output folder first")
    args = parser.parse_args()

    out: Path = args.out
    if args.clean and out.exists():
        shutil.rmtree(out)
    out.mkdir(parents=True, exist_ok=True)

    BODY_FONT, BOLD_FONT = _register_fonts()

    written = 0
    for report in REPORTS:
        name = f"{report['code']}.{report['format']}"
        path = out / name
        if report["format"] == "docx":
            write_docx(path, report["pages"])
        else:
            write_pdf(path, report["pages"])
        written += 1
        print(f"  {name:34} {report['expect']}")

    # A byte-identical copy under a different name, for the duplicate scan.
    source = out / "2026-SYN-e-NVH-01.pdf"
    duplicate = out / "2026-SYN-e-NVH-01_kopya.pdf"
    shutil.copyfile(source, duplicate)
    written += 1
    print(f"  {duplicate.name:34} identical bytes to NVH-01 - duplicate detection")

    catalog_path = out / "katalog_SYN.csv"
    with catalog_path.open("w", encoding="utf-8-sig", newline="") as handle:
        csv.writer(handle, delimiter=";").writerows(build_catalog_rows())
    print(f"  {catalog_path.name:34} {len(REPORTS)} catalog rows")

    print(f"\n{written} report files + 1 catalog written to {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
