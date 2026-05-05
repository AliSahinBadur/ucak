from __future__ import annotations

from pathlib import Path

from ..schemas import ParsedSection


def parse_docx(file_path: str | Path) -> list[ParsedSection]:
    """Extract headings, paragraphs, and simple table text from DOCX files."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX parsing requires the 'python-docx' package. Install project dependencies first."
        ) from exc

    path = Path(file_path)
    document = Document(str(path))
    sections: list[ParsedSection] = []
    current_title: str | None = None
    buffer: list[str] = []
    logical_page = 1

    def flush() -> None:
        nonlocal logical_page, buffer
        text = "\n".join(part for part in buffer if part.strip()).strip()
        if text:
            sections.append(
                ParsedSection(
                    page_number=logical_page,
                    raw_text=text,
                    section_title=current_title,
                )
            )
            logical_page += 1
        buffer = []

    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue

        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            flush()
            current_title = text
            buffer.append(text)
            continue

        buffer.append(text)

    table_lines = _extract_tables(document)
    if table_lines:
        buffer.extend(table_lines)

    flush()
    return sections


def _extract_tables(document) -> list[str]:
    lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            cleaned = [cell for cell in cells if cell]
            if cleaned:
                lines.append(" | ".join(cleaned))
    return lines
